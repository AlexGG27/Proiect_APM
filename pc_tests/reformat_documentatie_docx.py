from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def get_style(doc: Document, name: str):
    for style in doc.styles:
        if style.name and style.name.lower() == name.lower():
            return style
    raise KeyError(name)


def get_or_add_paragraph_style(doc: Document, name: str):
    for style in doc.styles:
        if style.name and style.name.lower() == name.lower():
            return style
    return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_paragraph_shading(paragraph, fill: str = "F5F5F5"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def clear_paragraph_shading(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is not None:
        p_pr.remove(shd)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def heading_level(text: str):
    text = text.strip()
    if re.match(r"^\d+\.\d+\.\d+\.\s+\S", text):
        return 3
    if re.match(r"^\d+\.\d+\.\s+\S", text):
        return 2
    if re.match(r"^\d+\.\s+\S", text):
        return 1
    return None


def is_caption(text: str) -> bool:
    text = text.strip()
    return bool(re.match(r"^(Figura|Fig\.|Tabelul|Tabel)\b", text, re.IGNORECASE))


def is_code_like(text: str) -> bool:
    text = text.rstrip()
    if not text:
        return False

    patterns = [
        r"^\s*#include\b",
        r"^\s*#define\b",
        r"^\s*\.SECTION\b",
        r"^\s*\.var\b",
        r"^\s*\.global\b",
        r"^\s*\.extern\b",
        r"^\s*(typedef|static|void|uint|int|switch|case|default:|return|while|for)\b",
        r"^\s*(TIM_|RCC_|GPIO_|HAL_)\w+",
        r"^\s*(ax0|ax1|ay0|ay1|ar|mr|mf|mx0|mx1|my0|my1|sr|si|tx0|none)\b",
        r"^\s*(dm|pm)\(",
        r"^\s*[A-Za-z_][A-Za-z0-9_]*:$",
        r"^\s*(jump|rti|ena|dis|imask|icntl|mstat|idle|nop|modify)\b",
        r"^\s*/\*",
        r"^\s*\*",
        r"^\s*\*/",
        r"^\s*{",
        r"^\s*}",
        r"^\s*\|",
    ]

    if any(re.match(pattern, text) for pattern in patterns):
        return True

    if text in {"AGC", "ALE", "MF"}:
        return True

    return False


def is_short_label(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    if is_caption(text) or is_code_like(text) or heading_level(text):
        return False
    return text.endswith(":") and len(text) <= 90


def next_non_empty_text(paragraphs, start_idx: int) -> str:
    for para in paragraphs[start_idx + 1:]:
        text = para.text.strip()
        if text:
            return text
    return ""


def is_technical_subheading(text: str, next_text: str) -> bool:
    text = text.strip()
    next_text = next_text.strip()

    if not text:
        return False
    if heading_level(text) or is_caption(text) or is_code_like(text):
        return False
    if text.endswith(":") or text.endswith(".") or len(text) > 120:
        return False

    looks_like_title = (
        " - " in text
        or "(" in text
        or text.lower().startswith(
            (
                "callback",
                "configurare",
                "citire",
                "scriere",
                "procesare",
                "decodare",
                "initializare",
                "selectie",
                "transmitere",
            )
        )
    )

    return looks_like_title and is_code_like(next_text)


def setup_styles(doc: Document):
    normal = get_style(doc, "normal")
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)

    heading1 = get_style(doc, "Heading 1")
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True

    heading2 = get_style(doc, "Heading 2")
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.keep_with_next = True

    heading3 = get_style(doc, "Heading 3")
    heading3.font.name = "Times New Roman"
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.italic = False
    heading3.font.color.rgb = RGBColor(0, 0, 0)
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(3)
    heading3.paragraph_format.keep_with_next = True

    caption = get_or_add_paragraph_style(doc, "LegendaP2")
    caption.base_style = normal
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    code = get_or_add_paragraph_style(doc, "CodP2")
    code.font.name = "Consolas"
    code.font.size = Pt(9.5)
    code.paragraph_format.left_indent = Cm(0.6)
    code.paragraph_format.right_indent = Cm(0.2)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def format_title_page(paragraphs):
    non_empty = [(i, p) for i, p in enumerate(paragraphs) if p.text.strip()]
    if not non_empty:
        return

    for i, para in non_empty:
        text = para.text.strip()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(3)
        clear_paragraph_shading(para)

        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        size = Pt(12)
        bold = False
        italic = False

        if i == non_empty[0][0]:
            size = Pt(14)
            bold = True
        elif text.startswith("APLICA"):
            size = Pt(14)
            bold = True
        elif "SISTEM DE PROCESARE DIGITALA" in text.upper() or "DSP (AGC, ALE, MF)" in text.upper():
            size = Pt(18)
            bold = True
        elif text.startswith("ECHIPA"):
            size = Pt(14)
            bold = True
            para.paragraph_format.space_after = Pt(8)
        elif text.startswith("Studen") or text.startswith("Profesor coordonator"):
            size = Pt(12)
            bold = True
        elif len(text) <= 40:
            size = Pt(12)
        else:
            size = Pt(12)
            italic = False

        for run in para.runs:
            run.font.size = size
            run.bold = bold
            run.italic = italic


def format_document(doc: Document):
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.start_type = WD_SECTION_START.CONTINUOUS

    paragraphs = doc.paragraphs

    first_heading_idx = None
    for idx, para in enumerate(paragraphs):
        txt = para.text.strip()
        if heading_level(txt) == 1 or para.style.name.lower() == "heading 1":
            first_heading_idx = idx
            break

    if first_heading_idx and first_heading_idx > 0:
        format_title_page(paragraphs[:first_heading_idx])

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        original_style_name = para.style.name.lower() if para.style.name else ""
        next_text = next_non_empty_text(paragraphs, i)

        if text == "\\":
            para.text = ""
            text = ""

        if not text:
            para.paragraph_format.space_after = Pt(0)
            clear_paragraph_shading(para)
            continue

        if i < (first_heading_idx or 0):
            continue

        if original_style_name == "heading 1":
            para.style = get_style(doc, "Heading 1")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue
        if original_style_name == "heading 2":
            para.style = get_style(doc, "Heading 2")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue
        if original_style_name == "heading 3":
            para.style = get_style(doc, "Heading 3")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue

        level = heading_level(text)
        if level == 1:
            para.style = get_style(doc, "Heading 1")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue
        if level == 2:
            para.style = get_style(doc, "Heading 2")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue
        if level == 3:
            para.style = get_style(doc, "Heading 3")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue

        if text in {"AGC", "ALE", "MF"}:
            para.style = get_style(doc, "Heading 3")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue

        if is_technical_subheading(text, next_text):
            para.style = get_style(doc, "Heading 3")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_shading(para)
            continue

        if is_caption(text):
            para.style = get_or_add_paragraph_style(doc, "LegendaP2")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_shading(para)
            continue

        if has_drawing(para):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_shading(para)
            continue

        if is_code_like(text):
            para.style = get_or_add_paragraph_style(doc, "CodP2")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_shading(para)
            continue

        para.style = get_style(doc, "normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        clear_paragraph_shading(para)

        if is_short_label(text):
            for run in para.runs:
                run.bold = True
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(2)

    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    para.style = get_style(doc, "normal")
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if r_idx == 0:
                        for run in para.runs:
                            run.bold = True


def main():
    if len(sys.argv) < 2:
        print("Usage: python reformat_documentatie_docx.py <input.docx> [output.docx]")
        raise SystemExit(1)

    input_path = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = input_path.with_name(input_path.stem + "_reformatata" + input_path.suffix)

    doc = Document(str(input_path))
    format_document(doc)
    doc.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
